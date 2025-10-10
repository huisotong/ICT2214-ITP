from flask import Blueprint, jsonify, request, Response, stream_with_context
from flask import current_app
from app.models.module_assignment import ModuleAssignment
from app.models.module import Module
from app.models.users import User
from app.models.chat_history import ChatHistory
from app.models.chat_message import ChatMessage
from app.models.chatbot_settings import ChatbotSettings
from app.models.agent import Agent
from app.db import db
from pathlib import Path
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct, PointIdsList
import os
import uuid
from dotenv import load_dotenv
from docx import Document as DOCXDocument
import pdfplumber
import openpyxl
from pptx import Presentation as PPTXDocument
from io import BytesIO
import traceback
from datetime import datetime
import requests
from queue import Queue
from threading import Thread
import json
from langchain.callbacks.base import BaseCallbackHandler
from langchain.schema import HumanMessage
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.callbacks import get_openai_callback

chatbot_bp = Blueprint('chatbot', __name__)
UPLOAD_FOLDER = Path("uploads")
UPLOAD_FOLDER.mkdir(exist_ok=True)

load_dotenv()

def get_qdrant_client():
    return QdrantClient(
        url=os.getenv("QDRANT_HOST"),
        api_key=os.getenv("QDRANT_API_KEY"),
        prefer_grpc=False,
        https=True,
        timeout=10.0,
        check_compatibility=False
    )

def internet_search_results(query, num_results=5):
    api_key = os.getenv("GOOGLE_API_KEY")
    cse_id = os.getenv("GOOGLE_CSE_ID")
    if not api_key or not cse_id:
        return ""
    try:
        response = requests.get(
            "https://www.googleapis.com/customsearch/v1",
            params={"key": api_key, "cx": cse_id, "q": query, "num": num_results},
            timeout=10
        )
        data = response.json()
        items = data.get("items", [])
        results = []
        for item in items:
            snippet = item.get("snippet", "")
            title = item.get("title", "")
            if snippet:
                results.append(f"{title}: {snippet}")
        return "\n".join(results)
    except Exception as e:
        print(f"Internet search failed: {e}")
        return ""

# Function to extract text
def extract_text_from_file(file):
    ext = file.filename.lower().split('.')[-1]
    content = ""

    # DOCX (Word)
    if ext == "docx":
        docx_io = BytesIO(file.read())
        doc = DOCXDocument(docx_io)

        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        content += "\n".join(paragraphs)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content += "\n" + "\t".join(row_text)

    # PDF
    elif ext == "pdf":
        pdf_io = BytesIO(file.read())
        with pdfplumber.open(pdf_io) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"

                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = [cell.strip() if cell else "" for cell in row]
                        content += "\n" + "\t".join(row_text)

    # PPTX (PowerPoint)
    elif ext == "pptx":
        pptx_io = BytesIO(file.read())
        prs = PPTXDocument(pptx_io)

        for slide in prs.slides:
            for shape in slide.shapes:
                if hasattr(shape, "text") and shape.text.strip():
                    content += shape.text.strip() + "\n"

                if shape.has_table:
                    for row in shape.table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        content += "\n" + "\t".join(row_text)

    # XLSX (Excel)
    elif ext in ["xlsx", "xls"]:
        excel_io = BytesIO(file.read())
        wb = openpyxl.load_workbook(excel_io, data_only=True)

        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(row_text):
                    content += "\n" + "\t".join(row_text)

    # Unsupported
    else:
        raise ValueError("Unsupported file format")

    return content.strip()

# Function to tag document from qdrant
def tag_document_to_qdrant(module_id: str, file_content: str, filename: str):
    # Step 1: Initialize embedding model and text splitter
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-base-en-v1.5")
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)

    # Step 2: Split file content into manageable chunks
    chunks = splitter.split_text(file_content)

    # Step 3: Wrap each chunk into a LangChain Document, attach filename as metadata
    documents = [
        Document(page_content=chunk, metadata={"filename": filename})
        for chunk in chunks
    ]

    # Step 4: Connect to Qdrant vector database
    client = get_qdrant_client()
    print(f"Connecting to Qdrant at: {os.getenv('QDRANT_HOST')}")
    collection_name = f"module_{module_id}"

    # Step 5: Create collection if it doesn't exist yet
    if collection_name not in [c.name for c in client.get_collections().collections]:
        dummy_vector = embeddings.embed_documents(["test"])[0]
        client.recreate_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=len(dummy_vector), distance=Distance.COSINE)
        )

    # Step 6: Add documents to the vector store
    print(f"Uploading {len(documents)} documents to Qdrant...")
    vectorstore = QdrantVectorStore(
        client=client,
        collection_name=collection_name,
        embedding=embeddings
    )
    vectorstore.add_documents(documents)
    print("Tagging completed successfully.")

    # Step 7: Return list of filenames tagged (same repeated filename for each chunk)
    return [str(doc.metadata["filename"]) for doc in documents]


# Function to untag document from qdrant
def untag_document_from_qdrant(module_id: str, filename: str):
    client = get_qdrant_client()
    collection_name = f"module_{module_id}"

    print(f"Removing all points for file '{filename}' from collection '{collection_name}'...")

    try:
        # Step 1: Retrieve all vectors and their payloads from the specified collection
        scroll_result = client.scroll(
            collection_name=collection_name,
            with_payload=True,
            limit=1000  # May increase if you expect very large files
        )
    except Exception as e:
        print(f"Failed to scroll Qdrant collection: {e}")
        return

    all_docs = scroll_result[0]
    matching_ids = []

    # Step 2: Check for filename in both flat and nested payload structures
    for doc in all_docs:
        filename_flat = doc.payload.get("filename")
        filename_nested = doc.payload.get("metadata", {}).get("filename")

        if filename_flat == filename:
            print(f"Matched flat filename: {filename_flat}")
            matching_ids.append(str(doc.id))
        elif filename_nested == filename:
            print(f"Matched nested filename: {filename_nested}")
            matching_ids.append(str(doc.id))

    if not matching_ids:
        print("⚠️ No points found with that filename.")
        return

    # Step 3: Delete all points that matched the filename
    result = client.delete(
        collection_name=collection_name,
        points_selector=PointIdsList(points=matching_ids)
    )

    print(f"Deleted {len(matching_ids)} points. Qdrant response:", result)


@chatbot_bp.route('/get-model-settings/<module_id>', methods=['GET'])
def get_model_settings(module_id):
    try:
        # Get LLM settings from sql
        settings = ChatbotSettings.query.filter_by(moduleID=module_id).first()
        if not settings:
            return jsonify({'error': 'Settings not found'}), 404

        # Initialize Qdrant client (need change depending on setup)
        client = get_qdrant_client()
        collection_name = f"module_{module_id}"
        filenames = []
        try:
            # Get all documents from the module's collection
            documents = client.scroll(collection_name=collection_name, with_payload=True, limit=1000)[0]
             # Extract filenames from flat or nested metadata
            for doc in documents:
                flat_name = doc.payload.get("filename")
                nested_name = doc.payload.get("metadata", {}).get("filename")
                name = flat_name or nested_name

                if name:
                    name = name.strip()
                    if name not in filenames:
                        filenames.append(name)
                    print(f"Found filename: {name}")
                else:
                    print(f"Skipped point without filename. Point ID: {doc.id}")


        except Exception as e:
            print(f"Error fetching from Qdrant: {str(e)}")
            filenames = []

        return jsonify({
            'settings': {
                'model': settings.model,
                'temperature': settings.temperature,
                'systemPrompt': settings.system_prompt,
                'maxTokens': settings.max_tokens,
            },
            "documents": [{"id": name, "name": name} for name in filenames]
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@chatbot_bp.route('/save-model-settings', methods=['PUT'])
def save_model_settings():
    try:
        data = request.json
        module_id = data.get('moduleID')
        
        # Save chatbot settings to MySQL
        settings = ChatbotSettings.query.filter_by(moduleID=module_id).first()
        if not settings:
            settings = ChatbotSettings(moduleID=module_id)
        
        settings.model = data.get('model')
        settings.temperature = data.get('temperature')
        settings.system_prompt = data.get('systemPrompt')
        settings.max_tokens = data.get('maxTokens')
        
        db.session.add(settings)
        db.session.commit()
    

        return jsonify({
            "status": "success",
            "message": "Settings saved successfully"
        }), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500
    
# Tag Document Route
@chatbot_bp.route('/tag-document', methods=['POST'])
def tag_document():
    try:
        print("Received request to /tag-document")
        module_id = request.form.get("moduleID")
        file = request.files.get("file")

        if not module_id or not file:
            return jsonify({"error": "Missing moduleID or file"}), 400

        # Extract text from any supported file
        content = extract_text_from_file(file)

        if not content.strip():
            return jsonify({"error": "No content extracted from file"}), 400

         # Step 2: Tag content into Qdrant via embedding and splitting
        part_filenames = tag_document_to_qdrant(module_id, content, file.filename)

        return jsonify({
            "status": "success",
            "filename": file.filename,
            "num_chunks": len(part_filenames),
            "chunks": part_filenames
        }), 200

    except Exception as e:
        print("EXCEPTION in /tag-document route:")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    

# Untag Document Route
@chatbot_bp.route('/untag-document', methods=['POST'])
def untag_document():
    try:
        print("Received request to /untag-document")
        data = request.get_json()
        module_id = data.get("moduleID")
        filename = data.get("docID")  # Must match the filename stored in metadata

        if not module_id or not filename:
            return jsonify({"error": "Missing moduleID or filename"}), 400

        # Step 1: Call helper to remove all matching Qdrant chunks
        untag_document_from_qdrant(module_id, filename)

        return jsonify({"status": "success"}), 200

    except Exception as e:
        print("EXCEPTION in /untag-document route:", str(e))
        return jsonify({"error": str(e)}), 500


# Sends user's message and streams chatbot's response
@chatbot_bp.route('/send-message', methods=['POST'])
def send_message():
    try:
        data = request.get_json()

        chat_id = data.get("chat_id")
        user_message = data.get("message")
        module_id = data.get("module_id")
        agent_id = data.get("agent_id")
        model_override = data.get("model")
        user_id = data.get("user_id")
        internet_search = data.get("internet_search", False)

        # 1) Validate + assignment/agent
        if not user_message or not user_id or not module_id:
            return jsonify({"error": "module_id, message, and user_id are required"}), 400

        if module_id and agent_id:
            return jsonify({"error": "Cannot specify both module_id and agent_id"}), 400

        if not module_id and not agent_id:
            return jsonify({"error": "Either module_id or agent_id is required"}), 400

        assignment = None
        if module_id:
            assignment = ModuleAssignment.query.filter_by(userID=user_id, moduleID=str(module_id)).first()
            if not assignment:
                return jsonify({"error": "No assignment found for the given user and module"}), 404

            # Negative credit gate for module-based chats
            if assignment.studentCredits is not None and assignment.studentCredits < 0:
                return jsonify({
                    "error": "Insufficient credits",
                    "message": "You have negative credits and cannot submit new prompts. Please request additional credits from your instructor.",
                    "current_credits": assignment.studentCredits
                }), 403

        # 2) Chat session (create or fetch)
        if not chat_id:
            if module_id:
                new_chat = ChatHistory(
                    assignmentID=assignment.assignmentID,
                    chatlog="",   # You store the generated title here
                    dateStarted=datetime.utcnow()
                )
            else:
                # For agent-based chats, create a chat session with agentID
                new_chat = ChatHistory(
                    assignmentID=None,  # No assignment for agent chats
                    agentID=agent_id,   # Link to the agent
                    chatlog="",   # You store the generated title here
                    dateStarted=datetime.utcnow()
                )
            db.session.add(new_chat)
            db.session.commit()
            chat_id = new_chat.historyID
            chat_session = new_chat
        else:
            chat_session = ChatHistory.query.filter_by(historyID=chat_id).first()
            if not chat_session:
                return jsonify({"error": "Chat session not found"}), 404

        # 3) Settings (+ optional model override)
        if module_id:
            settings = ChatbotSettings.query.filter_by(moduleID=str(module_id)).first()
            if not settings:
                return jsonify({"error": "Model settings not found for this module"}), 404
            if model_override:
                settings.model = model_override
        else:
            # For agent-based chats, get settings from the agent
            agent = Agent.query.filter_by(agentID=agent_id).first()
            if not agent:
                return jsonify({"error": "Agent not found"}), 404
            
            # Create a settings-like object from agent data
            class AgentSettings:
                def __init__(self, agent):
                    self.model = agent.model
                    self.temperature = agent.temperature
                    self.system_prompt = agent.system_prompt
                    self.max_tokens = agent.max_tokens
            
            settings = AgentSettings(agent)
            if model_override:
                settings.model = model_override

        # 4) Pricing lookup (OpenRouter) — only for module-based chats
        prompt_price = 0
        completion_price = 0
        
        if module_id:  # Only calculate costs for module-based chats
            try:
                models_response = requests.get("https://openrouter.ai/api/v1/models")
                models_response.raise_for_status()
                models_data = models_response.json().get("data", [])
            except requests.exceptions.RequestException as e:
                print(f"Could not fetch model pricing from OpenRouter: {e}")
                return jsonify({"error": "Could not fetch model pricing. Please try again."}), 500

            model_info = next((m for m in models_data if m.get("id") == settings.model), None)
            if not model_info:
                return jsonify({"error": f"Could not find pricing information for model: {settings.model}"}), 500
            pricing = model_info.get("pricing", {})
            prompt_price = float(pricing.get("prompt", 0))
            completion_price = float(pricing.get("completion", 0))

        # 5) System context
        system_context = f"System Context: {settings.system_prompt}\n\n" if settings.system_prompt else ""

        # 6) Build conversation history (your schema uses chatID on ChatMessage)
        previous_messages = (
            db.session.query(ChatMessage)
            .filter_by(chatID=chat_id)
            .order_by(ChatMessage.timestamp.asc())
            .all()
        )
        conversation_history = []
        current_pair = {"user": None, "ai": None}
        for msg in previous_messages:
            if msg.sender == 'user':
                if current_pair["user"] or current_pair["ai"]:
                    conversation_history.append((current_pair["user"] or "", current_pair["ai"] or ""))
                    current_pair = {"user": None, "ai": None}
                current_pair["user"] = msg.content
            elif msg.sender == 'ai':
                current_pair["ai"] = msg.content
                conversation_history.append((current_pair["user"] or "", current_pair["ai"] or ""))
                current_pair = {"user": None, "ai": None}

        # 7) Generate a chat title on first message (same logic)
        if not previous_messages:
            title_prompt = (
                f"Provide one short, descriptive chat title for the following conversation. "
                f"Return only the title, without numbering or additional commentary: {user_message}"
            )
            llm_for_title = ChatOpenAI(
                model_name=settings.model,
                temperature=0.7,
                max_tokens=20,
                openai_api_key=os.getenv("OPENROUTER_API_KEY"),
                openai_api_base=os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
            )
            title_response = llm_for_title.invoke([HumanMessage(content=title_prompt)])
            chat_title = title_response.content.strip()
            if chat_title.startswith('"') and chat_title.endswith('"'):
                chat_title = chat_title[1:-1].strip()
            elif chat_title.startswith("'") and chat_title.endswith("'"):
                chat_title = chat_title[1:-1].strip()
            chat_session.chatlog = chat_title
            db.session.commit()

        # 8) Qdrant: detect if documents exist (same)
        collection_name = f"module_{module_id}"
        client = get_qdrant_client()
        try:
            # Instead of fetching .scroll(), check if collection actually exists
            collections = [c.name for c in client.get_collections().collections]
            if collection_name in collections:
                documents = True  # Flag that Qdrant context is available
            else:
                documents = False
        except Exception as q_err:
            print(f"Error checking Qdrant collection existence: {str(q_err)}")
            documents = False


        # ===== Streaming setup =====
        class _QueueTokenHandler(BaseCallbackHandler):
            def __init__(self, q):
                self.q = q
            def on_llm_new_token(self, token: str, **kwargs):
                self.q.put({"type": "token", "data": token})
            def on_llm_end(self, response, **kwargs):
                self.q.put({"type": "lc_end"})

        def _sse(data: dict) -> bytes:
            return f"data: {json.dumps(data, ensure_ascii=False)}\n\n".encode("utf-8")

        q = Queue()
        cb = _QueueTokenHandler(q)

        streaming_llm = ChatOpenAI(
            model_name=settings.model,
            temperature=settings.temperature,
            max_tokens=settings.max_tokens,
            streaming=True,
            callbacks=[cb],
            openai_api_key=os.getenv("OPENROUTER_API_KEY"),
            openai_api_base=os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
        )

        accumulated = {"text": ""}
        usage = {"prompt": 0, "completion": 0}  # to collect tokens for costing

                # Producer runs the LLM chain and pushes tokens into the queue
        def _producer():
            try:
                with get_openai_callback() as cb_usage:
                    search_context = ""
                    if internet_search:
                        try:
                            search_context = internet_search_results(user_message)
                        except Exception as e:
                            print("Internet search failed:", e)
                            search_context = ""

                    # If the collection exists we built 'documents' flag earlier
                    if documents:
                        embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-base-en-v1.5")
                        vectorstore = QdrantVectorStore(client=client, collection_name=collection_name, embedding=embeddings)
                        retriever = vectorstore.as_retriever(search_kwargs={"k": 3, "score_threshold": 0.5})

                        # 1) Get top retrieved docs first (synchronous check)
                        retrieved_docs = retriever.get_relevant_documents(user_message) or []
                        print(f"Retrieved {len(retrieved_docs)} docs for query: {user_message}")

                        # 2) If no retrieved docs AND internet_search is False => immediately respond "outside scope"
                        if not retrieved_docs and not search_context:
                            print("No docs found, sending outside-scope message.")
                            outside_msg = "I'm sorry — this topic is outside this module's scope. Please enable Internet Search for an open-web answer."
                            # push one 'token' event first so frontend displays text
                            q.put({"type": "token", "data": outside_msg})
                            q.put({
                                "type": "done",
                                "final": outside_msg,
                                "chat_id": chat_id,
                                "chat_title": chat_session.chatlog,
                                "cost": 0
                            })
                            return
                        # 3) Build combined context: docs (if any) + web snippets (if any)
                        combined_context = ""
                        if retrieved_docs:
                            combined_context += "\n\n".join([d.page_content for d in retrieved_docs])
                        if search_context:
                            if combined_context:
                                combined_context += "\n\n"
                            combined_context += f"[Web Snippets]\n{search_context}"

                        # 4) Create a strict system instruction depending on internet_search
                        if search_context:
                            # allow mixing of docs + web
                            system_instr = (
                                "You are a teaching assistant for this course. Use the documents from the module first, "
                                "and also the provided web snippets. Answer accurately and concisely, and cite sources when helpful."
                            )
                        else:
                            # docs only -> restrict to docs
                            system_instr = (
                                "You are a teaching assistant for this course. Only answer using the provided module documents. "
                                "If the answer cannot be found in the documents, reply: "
                                "'I'm sorry, this topic is outside the module’s scope.' Do not use your own external knowledge."
                            )

                        prompt = (
                            f"{system_instr}\n\nContext:\n{combined_context}\n\nQuestion: {user_message}\nAI:"
                        )

                        # 5) Stream tokens from the LLM using the same streaming_llm and callbacks
                        streaming_llm.invoke([HumanMessage(content=prompt)])

                    else:
                        # No qdrant collection at all: fallback
                        # if internet_search is enabled => use web
                        if internet_search and search_context:
                            prompt_text = (
                                f"Use the following web snippets to answer the question. "
                                f"Web snippets:\n{search_context}\n\nQuestion: {user_message}\nAI:"
                            )
                            streaming_llm.invoke([HumanMessage(content=prompt_text)])
                        else:
                            # No docs and internet search disabled -> outside scope
                            outside_msg = "I'm sorry — this topic is outside this module's scope. Please enable Internet Search for an open-web answer."
                            q.put({"type": "done", "final": outside_msg, "chat_id": chat_id, "chat_title": chat_session.chatlog, "cost": 0})

                    # Track token usage after LLM returns
                    usage["prompt"] = cb_usage.prompt_tokens or 0
                    usage["completion"] = cb_usage.completion_tokens or 0

            except Exception as e:
                print("Exception inside _producer:", e)
                q.put({"type": "error", "message": str(e)})
            finally:
                q.put({"type": "end"})


        # Capture the current Flask app *while still inside the request context*
        flask_app = current_app._get_current_object()

        def thread_wrapper():
            # Push the captured app context into the new thread
            with flask_app.app_context():
                _producer()

        t = Thread(target=thread_wrapper, daemon=True)
        t.start()


        @stream_with_context
        def event_stream():
            # Start event
            yield _sse({"type": "start"})

            while True:
                item = q.get()
                if item["type"] == "token":
                    tok = item["data"]
                    accumulated["text"] += tok
                    yield _sse({"type": "token", "data": tok})
                elif item["type"] == "error":
                    yield _sse({"type": "error", "message": item.get("message", "unknown error")})
                    break
                elif item["type"] in ("lc_end", "end"):
                    # 9) Costing & credit deduction (only for module-based chats)
                    try:
                        cost = 0
                        if module_id:  # Only calculate costs for module-based chats
                            cost = (usage["prompt"] * prompt_price) + (usage["completion"] * completion_price)
                            # Deduct and save assignment
                            if assignment and assignment.studentCredits is not None:
                                assignment.studentCredits -= cost
                                db.session.add(assignment)

                        # Save both messages now
                        user_msg = ChatMessage(
                            chatID=chat_id,
                            sender="user",
                            content=user_message,
                            timestamp=datetime.utcnow()
                        )
                        bot_msg = ChatMessage(
                            chatID=chat_id,
                            sender="ai",
                            content=accumulated["text"],
                            timestamp=datetime.utcnow()
                        )
                        db.session.add(user_msg)
                        db.session.add(bot_msg)
                        db.session.commit()
                    except Exception as e:
                        db.session.rollback()
                        yield _sse({"type": "error", "message": f"Failed to save or deduct credits: {str(e)}"})
                        break

                    # 10) Final "done"
                    yield _sse({
                        "type": "done",
                        "chat_id": chat_id,
                        "chat_title": chat_session.chatlog,  # title stored in chatlog in your schema
                        "final": accumulated["text"],
                        "cost": cost
                    })
                    break

        headers = {
            "Content-Type": "text/event-stream; charset=utf-8",
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # if behind nginx
        }
        return Response(event_stream(), headers=headers)

    except Exception as e:
        traceback.print_exc()
        db.session.rollback()
        return jsonify({"error": str(e)}), 500



@chatbot_bp.route('/get-chat-history/<int:chat_id>', methods=['GET'])
def get_chat_history(chat_id):
    try:
        messages = (
            db.session.query(ChatMessage)
            .filter_by(chatID=chat_id)
            .order_by(ChatMessage.timestamp.asc())
            .all()
        )
        history = [{
            "sender": msg.sender,
            "content": msg.content,
            "timestamp": msg.timestamp.isoformat()
        } for msg in messages]

        return jsonify({"chat_history": history}), 200

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@chatbot_bp.route('/get-chat-history/<user_id>/<module_id>', methods=['GET'])
def get_chat_history_for_user_module(user_id, module_id):
    """
    Returns all chat sessions (with all columns) belonging to the assignment for a given user and module.
    """
    try:
        assignment = ModuleAssignment.query.filter_by(userID=user_id, moduleID=str(module_id)).first()
        if not assignment:
            return jsonify({"error": "No assignment found for the given user and module"}), 404

        chats = ChatHistory.query.filter_by(assignmentID=assignment.assignmentID) \
            .order_by(ChatHistory.dateStarted.desc()).all()
        chat_list = []
        for chat in chats:
            chat_list.append({
                "historyID": chat.historyID,
                "assignmentID": chat.assignmentID,
                "chatlog": chat.chatlog,
                "dateStarted": chat.dateStarted.isoformat()
            })
        return jsonify(chat_list), 200
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@chatbot_bp.route('/get-chat-message/<int:chat_id>', methods=['GET'])
def get_chat_message(chat_id):
    """
    Returns all messages (with all columns) for the given chat session.
    """
    try:
        messages = db.session.query(ChatMessage) \
            .filter_by(chatID=chat_id) \
            .order_by(ChatMessage.timestamp.asc()).all()
        msg_list = []
        for msg in messages:
            msg_list.append({
                "chatID": msg.chatID,
                "sender": msg.sender,
                "content": msg.content,
                "timestamp": msg.timestamp.isoformat()
            })
        return jsonify(msg_list), 200
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@chatbot_bp.route("/get-module-model/<string:module_id>", methods=["GET"])
def get_module_model(module_id):
    chatbot_settings = ChatbotSettings.query.filter_by(moduleID=str(module_id)).first()
    if not chatbot_settings:
        return jsonify({"error": "Chatbot settings not found for this module"}), 404

    return jsonify(
        {
            "model_name": chatbot_settings.model,
        }
    )